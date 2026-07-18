from io import BytesIO
from docx import Document
from edu_copilot.reports.report_data import ReportData

def generate_docx_report(data: ReportData, output_stream: BytesIO) -> None:
    """
    Generates a professional academic intervention report in DOCX format
    from a ReportData instance and writes it to the output stream.
    
    Args:
        data (ReportData): Encapsulated report data.
        output_stream (BytesIO): Stream to save DOCX binary to.
    """
    doc = Document()
    
    # 1. Main Heading
    doc.add_heading("Student Academic Intervention Report", level=0)
    
    # 2. Metadata Grid Table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = 'Light Shading Accent 1'
    
    r1_cells = meta_table.rows[0].cells
    r1_cells[0].text = f"Student Name: {data.student_name}"
    r1_cells[1].text = f"Student ID: {data.student_id}"
    
    r2_cells = meta_table.rows[1].cells
    r2_cells[0].text = f"Grade Level: {data.grade_level}"
    r2_cells[1].text = f"Review Status: {data.reviewed_status}"
    
    r3_cells = meta_table.rows[2].cells
    r3_cells[0].text = f"Assessed Risk Level: {data.predicted_risk}"
    r3_cells[1].text = f"Prediction Confidence: {data.confidence_score:.1%}"
    
    doc.add_paragraph() # spacing
    
    # 3. Section 1: Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(data.summary)
    
    # 4. Section 2: Diagnostics & Metrics Table
    doc.add_heading("2. Core Triggers & Diagnostics", level=1)
    doc.add_paragraph(data.reasoning)
    
    # Quantitative stats grid
    metrics_table = doc.add_table(rows=6, cols=2)
    metrics_table.style = 'Table Grid'
    
    metrics_mapping = [
        ("Academic Indicator", "Observed Value"),
        ("Cumulative GPA", f"{data.gpa:.2f}"),
        ("Attendance Rate", f"{data.attendance_rate:.1%}"),
        ("Weekly Study Commitment", f"{data.study_hours:.1f} hours"),
        ("Prior Exam Grade", f"{data.previous_grade:.1f}%"),
        ("Parental Involvement Level", data.parental_involvement)
    ]
    
    for idx, (metric, val) in enumerate(metrics_mapping):
        cells = metrics_table.rows[idx].cells
        cells[0].text = metric
        cells[1].text = val
        
    doc.add_paragraph() # spacing
    
    # 4.1. Multimodal Sub-scores Table
    doc.add_heading("Model Multimodal Sub-Scores", level=2)
    subscores_table = doc.add_table(rows=6, cols=2)
    subscores_table.style = 'Table Grid'
    
    cnn_mean = (data.cnn_legibility + data.cnn_correctness + data.cnn_completeness) / 3.0 if data.cnn_legibility is not None else None
    
    subscores_mapping = [
        ("Predictive Modality", "Sub-Score Value"),
        ("Tabular ANN Student Risk", f"{data.tabular_score:.1%}" if data.tabular_score is not None else "N/A"),
        ("Worksheet CNN Quality (Mean)", f"{cnn_mean:.1%} (Legibility: {data.cnn_legibility:.0%}, Correctness: {data.cnn_correctness:.0%}, Completeness: {data.cnn_completeness:.0%})" if cnn_mean is not None else "N/A"),
        ("Weekly Activity LSTM Risk", f"{data.timeseries_score:.1%}" if data.timeseries_score is not None else "N/A"),
        ("Oral Exam GRU Fluency", f"{data.audio_score:.1%}" if data.audio_score is not None else "N/A"),
        ("Fused Composite Recommendation Risk", f"{data.fused_score:.1%}" if data.fused_score is not None else "N/A")
    ]
    
    for idx, (metric, val) in enumerate(subscores_mapping):
        cells = subscores_table.rows[idx].cells
        cells[0].text = metric
        cells[1].text = val
        
    doc.add_paragraph() # spacing
    
    # 5. Section 3: Action Plan
    doc.add_heading("3. Actionable Intervention Plan", level=1)
    for rec in data.recommendations:
        doc.add_paragraph(rec, style='List Bullet')
        
    # 5.1. Section 4: Explainability Maps
    import os
    from docx.shared import Inches
    
    explain_header = False
    
    # Check and add Grad-CAM image
    if data.gradcam_image_path and os.path.exists(data.gradcam_image_path):
        if not explain_header:
            doc.add_heading("4. Multimodal Explainability Artifacts", level=1)
            doc.add_paragraph("Visual explanations of model attributions across modalities:")
            explain_header = True
        doc.add_paragraph("Worksheet CNN Grad-CAM Saliency:")
        doc.add_picture(data.gradcam_image_path, width=Inches(3.2))
        doc.add_paragraph()
            
    # Check and add time-series attention chart
    if data.timeseries_attention_path and os.path.exists(data.timeseries_attention_path):
        if not explain_header:
            doc.add_heading("4. Multimodal Explainability Artifacts", level=1)
            doc.add_paragraph("Visual explanations of model attributions across modalities:")
            explain_header = True
        doc.add_paragraph("LSTM Time Series Temporal Attentions:")
        doc.add_picture(data.timeseries_attention_path, width=Inches(4.2))
        doc.add_paragraph()
            
    # Check and add audio attention chart
    if data.audio_attention_path and os.path.exists(data.audio_attention_path):
        if not explain_header:
            doc.add_heading("4. Multimodal Explainability Artifacts", level=1)
            doc.add_paragraph("Visual explanations of model attributions across modalities:")
            explain_header = True
        doc.add_paragraph("GRU Oral Audio Response Temporal Saliency:")
        doc.add_picture(data.audio_attention_path, width=Inches(4.2))
        doc.add_paragraph()
        
    # Optional Human Audit Notes
    if data.reviewer_notes:
        doc.add_paragraph()
        doc.add_heading("Human Review Audit Notes", level=1)
        doc.add_paragraph(data.reviewer_notes)
        
    doc.save(output_stream)
